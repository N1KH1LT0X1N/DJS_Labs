from docx import Document
import sys

def read_docx(filepath):
    doc = Document(filepath)
    print(f"\n{'='*80}")
    print(f"Reading: {filepath}")
    print(f"{'='*80}\n")
    
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")
    print("\n--- PARAGRAPHS ---\n")
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"{para.text}")
    
    if doc.tables:
        print("\n--- TABLES ---\n")
        for table_idx, table in enumerate(doc.tables):
            print(f"\n=== Table {table_idx + 1} ===")
            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                print(" | ".join(row_data))

if __name__ == "__main__":
    read_docx('CMPM_Expt0.docx')
    print("\n\n")
    read_docx('CMPM_Expt1.docx')
